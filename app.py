from flask import Flask, render_template, request, session, redirect, url_for, send_file, jsonify
import math
from io import BytesIO
from docx import Document
import os # Added to access environment variables
from dotenv import load_dotenv # Added to load .env file

load_dotenv() # Load environment variables from .env file

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY') # Load secret key from environment

# Available battery and module sizes (Ah and Wp)
BATTERY_SIZES = [10, 20, 30, 40, 50, 65, 75, 100, 150, 200]

@app.route('/robots.txt')
def robots_txt():
    return send_file('static/robots.txt')

@app.route('/favicon.ico')
def favicon():
    return send_file('static/favicon.ico', mimetype='image/vnd.microsoft.icon')

# Determine system voltage based on daily energy demand
def determine_system_voltage(daily_energy_demand):
    # Updated to avoid gaps between 2580 and 2902
    if daily_energy_demand <= 537:
        return 12
    elif daily_energy_demand <= 2580:
        return 24
    else:
        return 48

@app.route('/', methods=['GET', 'POST'])
def home():
    if request.method == 'POST':
        try:
            # Capture all form data into session so it can be reused if an error occurs
            session['inputs'] = request.form.to_dict(flat=False)

            # Load Inputs
            rooms = request.form.getlist('room[]')
            loads = request.form.getlist('load[]')
            quantities = request.form.getlist('qty[]')
            powers = request.form.getlist('power[]')
            hours = request.form.getlist('hours[]')

            if not (quantities and powers and hours):
                return "Error: Load details must be provided."

            # Other Inputs (convert to float or int as needed)
            losses = float(request.form.get('losses', 0))
            autonomy_days = float(request.form.get('autonomy_days', 1))
            discharge_limit = float(request.form.get('discharge_limit', 50))
            battery_efficiency = float(request.form.get('battery_efficiency', 80))
            insolation = float(request.form.get('insolation', 5))
            tilt_angle = float(request.form.get('tilt_angle', 15))

            # --- Tilt Angle Influence on Insolation ---
            # Assumed ideal tilt for general year-round performance (degrees).
            # In a more advanced model, this could be based on latitude.
            ideal_tilt_angle = 30.0
            
            # K_factor_for_cos scales the deviation.
            # Chosen such that a 45-degree deviation from ideal might result in
            # cos(45 * 4/3) = cos(60 degrees) = 0.5 correction factor.
            K_factor_for_cos = 4.0 / 3.0
            
            deviation_from_ideal = abs(tilt_angle - ideal_tilt_angle)
            
            # Angle to be used in cosine function, scaled by K_factor
            scaled_deviation_for_cos_degrees = deviation_from_ideal * K_factor_for_cos
            
            # Cap the effective angle for cosine to 90 degrees (where cos is 0)
            # to prevent issues with very large deviations.
            if scaled_deviation_for_cos_degrees > 90.0:
                scaled_deviation_for_cos_degrees = 90.0
                
            tilt_correction_factor = math.cos(math.radians(scaled_deviation_for_cos_degrees))
            
            # Apply a floor to the correction factor (e.g., 0.1 for very poor, but not impossible, tilts)
            if tilt_correction_factor < 0.1:
                tilt_correction_factor = 0.1
            # Ensure factor does not exceed 1.0 (it shouldn't with cos if ideal_tilt_angle is the peak)
            if tilt_correction_factor > 1.0:
                tilt_correction_factor = 1.0

            effective_insolation = insolation * tilt_correction_factor
            # --- End Tilt Angle Influence ---

            inverter_efficiency = float(request.form.get('inverter_efficiency', 90)) / 100.0
            inverter_type = request.form.get('inverter_type', 'Pure Sine Wave')
            surge_factor = float(request.form.get('surge_factor', 2))
            adjustment_factor = float(request.form.get('adjustment_factor', 1.2))
            module_rating = float(request.form.get('module_rating', 300))

            # Convert lists to numerical types
            quantities = [int(q) for q in quantities]
            powers = [float(p) for p in powers]
            hours = [float(h) for h in hours]

            # Total Daily Energy (Wh)
            total_energy = sum(q * p * h for q, p, h in zip(quantities, powers, hours))

            # Total Daily Energy with Losses
            total_energy_with_losses = total_energy * (1 + losses / 100)

            # System Voltage
            system_voltage = determine_system_voltage(total_energy_with_losses)

            # Battery Sizing
            required_battery_capacity = (
                total_energy_with_losses * autonomy_days /
                (system_voltage * (discharge_limit / 100) * (battery_efficiency / 100))
            )
            chosen_battery_capacity = min(
                [b for b in BATTERY_SIZES if b >= required_battery_capacity],
                default=max(BATTERY_SIZES)
            )

            # Solar Array and Modules
            array_size = (total_energy_with_losses / (effective_insolation * 0.8)) * adjustment_factor
            num_modules = math.ceil(array_size / module_rating)

            # Inverter Sizing
            max_load_watt = sum(q * p for q, p in zip(quantities, powers))
            inverter_size = math.ceil((max_load_watt * surge_factor * 1.2) / inverter_efficiency)

            # Charge Controller
            controller_current = math.ceil((array_size / system_voltage) * 1.25)

            # Save Results in session
            session['results'] = {
                'total_energy': round(total_energy, 2),
                'total_energy_with_losses': round(total_energy_with_losses, 2),
                'system_voltage': system_voltage,
                'days_of_autonomy': autonomy_days,
                'required_battery_capacity': round(required_battery_capacity, 2),
                'chosen_battery_capacity': chosen_battery_capacity,
                'insolation': insolation, # Original user input for reference
                'tilt_angle': tilt_angle,
                'tilt_correction_factor': round(tilt_correction_factor, 3), # Added for transparency
                'effective_insolation': round(effective_insolation, 2), # Added for transparency
                'array_size': round(array_size, 2),
                'module_rating': module_rating,
                'num_modules': num_modules,
                'rounded_controller_current': controller_current,
                'inverter_size': inverter_size,
                'inverter_type': inverter_type,
                'inverter_efficiency': inverter_efficiency * 100,
                'surge_factor': surge_factor,
                'adjustment_factor': adjustment_factor
            }

            # Determine tilt warning
            tilt_factor = session['results']['tilt_correction_factor']
            if tilt_factor < 0.4:
                session['results']['tilt_warning_message'] = "Poor tilt angle, significantly impacting energy generation."
                session['results']['tilt_warning_color'] = "red"
            elif tilt_factor < 0.8:
                session['results']['tilt_warning_message'] = "Moderate tilt angle, some impact on energy generation."
                session['results']['tilt_warning_color'] = "amber"
            else:
                session['results']['tilt_warning_message'] = "Good tilt angle, optimal energy generation."
                session['results']['tilt_warning_color'] = "green"

            # Determine Effective Insolation warning
            effective_insolation_value = session['results']['effective_insolation']
            if effective_insolation_value < 2: # Low insolation
                session['results']['effective_insolation_message'] = "Low effective insolation (< 2 kWh/m²/day), may require a significantly larger solar array."
                session['results']['effective_insolation_color'] = "red"
            elif effective_insolation_value < 4: # Moderate insolation
                session['results']['effective_insolation_message'] = "Moderate effective insolation (2-4 kWh/m²/day)."
                session['results']['effective_insolation_color'] = "amber"
            else: # Good insolation
                session['results']['effective_insolation_message'] = "Good effective insolation (>= 4 kWh/m²/day) for energy generation."
                session['results']['effective_insolation_color'] = "green"

            return redirect(url_for('results'))

        except Exception as e:
            return f"Error: {str(e)}"

    # Handle GET Request
    # Unpack session inputs if available
    inputs = {}
    raw_inputs_from_session = session.get('inputs', {}) # Use .get for safety

    # Fields that are expected to be arrays in the form
    ARRAY_FIELDS = {'room[]', 'load[]', 'qty[]', 'power[]', 'hours[]'}

    for key, val_list in raw_inputs_from_session.items():
        if key in ARRAY_FIELDS:
            # Always keep these as lists, even if they contain a single item
            inputs[key] = val_list
        elif val_list and len(val_list) == 1: # Check if val_list is not empty before accessing len
            # For non-array fields, flatten if it's a single-item list
            inputs[key] = val_list[0]
        else:
            # For non-array fields with multiple items (shouldn't happen for scalars) or empty lists
            inputs[key] = val_list

    return render_template('index.html', inputs=inputs)

@app.route('/customize')
def customize_components_page():
    required_ah = request.args.get('required_ah', None)
    system_voltage = request.args.get('system_voltage', None)
    total_energy_with_losses = request.args.get('total_energy_with_losses', None)
    effective_insolation = request.args.get('effective_insolation', None)
    adjustment_factor = request.args.get('adjustment_factor', None)
    return render_template('customize_components.html',
                           required_ah=required_ah,
                           system_voltage=system_voltage,
                           total_energy_with_losses=total_energy_with_losses,
                           effective_insolation=effective_insolation,
                           adjustment_factor=adjustment_factor)

@app.route('/api/save_customizer_results', methods=['POST'])
def save_customizer_results_api():
    try:
        data = request.get_json()
        session['customizer_report_data'] = data
        return jsonify(success=True, message="Customizer results saved to session.")
    except Exception as e:
        # Log the exception for debugging
        print(f"Error in save_customizer_results_api: {str(e)}")
        return jsonify(success=False, message=str(e)), 500

@app.route('/download_customizer_report')
def download_customizer_report_route():
    results = session.get('customizer_report_data', None)
    if not results:
        return redirect(url_for('customize_components_page', error="No data available for report"))

    try:
        doc = Document()
        doc.add_heading('Solar Component Customizer Report', 0)

        def add_para(label, value, unit=""):
            if value is not None and str(value).strip() != '' and value != 'N/A': # Check for None, empty string and 'N/A' string
                doc.add_paragraph(f"{label}: {value}{unit}")
            else:
                doc.add_paragraph(f"{label}: Not calculated or N/A")

        doc.add_heading('System & Battery Inputs', level=1)
        add_para('System Voltage', results.get('Vsys'), "V")
        add_para('Battery Type', results.get('batType'))
        add_para('System Efficiency', results.get('effPercent'), "%")
        add_para('Depth of Discharge (DoD)', results.get('dodPercent'), "%")
        if results.get('Wh') is not None and results.get('autonomyDays') is not None:
            add_para('Daily Energy Need', results.get('Wh'), "Wh")
            add_para('Days of Autonomy', results.get('autonomyDays'), " days")
        elif results.get('extAh') is not None:
            add_para('Externally Provided Required Ah', results.get('extAh'), "Ah")

        doc.add_heading('Battery Bank Configuration', level=1)
        vsys_val = results.get('Vsys', '')
        add_para('Required Total Battery Capacity', results.get('reqAh'), f" Ah at {vsys_val}V")
        add_para('Individual Battery Voltage', results.get('Vbat'), "V")
        add_para('Individual Battery Capacity', results.get('AhBat'), "Ah (Nominal)")
        add_para('Total Batteries Needed', results.get('totalBattCount'), " units")
        wiring_text = f"{results.get('numStringsParallel', 'N/A')} parallel string(s), each with {results.get('battInSeries', 'N/A')} batteries in series."
        doc.add_paragraph(f"Wiring: {wiring_text}")

        if results.get('panelCalcPerformed'):
            doc.add_heading('PV Array & Charge Controller', level=1)
            add_para('Calculated Total Array Size (Ideal)', results.get('array_size_wp'), " Wp")
            doc.add_paragraph("") 
            p = doc.add_paragraph()
            p.add_run('Custom Panel Configuration:').bold = True
            add_para('  Individual Panel Power', results.get('Ppanel'), "Wp")
            add_para('  Individual Panel Voltage (Vmp)', results.get('Vpanel'), "V")
            
            num_panels_custom_val = results.get('num_panels_custom', 'N/A')
            raw_num_panels_val = results.get('raw_num_panels', 'N/A')
            if num_panels_custom_val != 'N/A' and raw_num_panels_val != 'N/A':
                doc.add_paragraph(f"  Calculated Number of Panels (Ideal): {num_panels_custom_val} (rounded up from {raw_num_panels_val})")
            elif num_panels_custom_val != 'N/A':
                doc.add_paragraph(f"  Calculated Number of Panels (Ideal): {num_panels_custom_val}")
            else:
                add_para('  Calculated Number of Panels (Ideal)', 'N/A')

            panel_wiring_text = f"{results.get('numPanelStringsParallel', 'N/A')} parallel string(s) of {results.get('panelsInSeries', 'N/A')} panels in series."
            doc.add_paragraph(f"  Panel Wiring Configuration: {panel_wiring_text}")
            add_para('  Total Panels Used in Configuration', results.get('totalPanelsConfigured'), " units")
            add_para('  Total Actual Panel Power in Configuration', results.get('totalActualPanelPower'), " Wp")
            doc.add_paragraph("") 
            p_cc = doc.add_paragraph()
            p_cc.add_run('Charge Controller Sizing:').bold = True
            add_para('  Recommended Charge Controller Size', results.get('recCC'), " A")
            
            cc_basis_text = f"  (Calculated based on the total actual panel power of {results.get('totalActualPanelPower', 'N/A')} Wp at {results.get('Vsys', 'N/A')}V system voltage, with a safety margin.)"
            doc.add_paragraph(cc_basis_text)
        else:
            doc.add_heading('PV Array & Charge Controller', level=1)
            doc.add_paragraph("Panel calculations were not performed or were incomplete due to missing inputs.")

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name='Solar_Component_Customizer_Report.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error in download_customizer_report_route: {str(e)}")
        # Maybe redirect with an error message to the customizer page
        return redirect(url_for('customize_components_page', error="Error generating report: " + str(e)))

@app.route('/results')
def results():
    results = session.get('results', {})
    if not results:
        return redirect(url_for('home'))
    return render_template('results.html', results=results)

@app.route('/download_report')
def download_report():
    results = session.get('results', {})
    if not results:
        return redirect(url_for('home'))

    # Generate Word doc
    doc = Document()
    doc.add_heading(f"System Sizing Results ({results['system_voltage']} V)", 0)

    doc.add_heading("System Overview", level=1)
    doc.add_paragraph(f"Total Daily Energy (Wh): {results['total_energy']}")
    doc.add_paragraph(f"Total Daily Energy with Losses (Wh): {results['total_energy_with_losses']}")
    doc.add_paragraph(f"System Voltage (V): {results['system_voltage']}")
    doc.add_paragraph(f"Days of Autonomy: {results['days_of_autonomy']}")
    doc.add_paragraph(f"Required Battery Capacity (Ah): {results['required_battery_capacity']}")
    doc.add_paragraph(f"Chosen Battery Capacity (Ah): {results['chosen_battery_capacity']}")
    doc.add_paragraph(f"Tilt Angle (°): {results['tilt_angle']}")
    # Display the original insolation, tilt correction factor, and effective insolation
    doc.add_paragraph(f"Original Insolation (kWh/m²/day): {results.get('insolation', 'N/A')}")
    doc.add_paragraph(f"Tilt Correction Factor: {results.get('tilt_correction_factor', 'N/A')}")
    doc.add_paragraph(f"Effective Insolation (kWh/m²/day): {results.get('effective_insolation', 'N/A')}")

    doc.add_heading("Solar Array", level=1)
    doc.add_paragraph(f"Array Size (Wp): {results['array_size']}")
    doc.add_paragraph(f"Module Rating (Wp): {results['module_rating']}")
    doc.add_paragraph(f"Number of Modules: {results['num_modules']}")
    doc.add_paragraph(f"Adjustment Factor: {results['adjustment_factor']}")

    doc.add_heading("Inverter", level=1)
    doc.add_paragraph(f"Inverter Type: {results['inverter_type']}")
    doc.add_paragraph(f"Inverter Efficiency (%): {results['inverter_efficiency']}")
    doc.add_paragraph(f"Inverter Size (W): {results['inverter_size']}")
    doc.add_paragraph(f"Surge Factor: {results['surge_factor']}")

    doc.add_heading("Charge Controller", level=1)
    doc.add_paragraph(f"Controller Current (A): {results['rounded_controller_current']}")

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name='Solar_Sizing_Report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/readme')
def readme():
    return render_template('README.html')

@app.route('/chat_log')
def chat_log_page():
    return render_template('chat_log_readme.html')

if __name__ == '__main__':
    print("Starting Flask app in development mode for debugging...")
    app.run(host='0.0.0.0', port=8080, debug=True)
